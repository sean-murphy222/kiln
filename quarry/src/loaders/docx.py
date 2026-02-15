"""
DOCX document loader using python-docx.

Extracts text, tables, and metadata from Word documents.
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any, ClassVar

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from chonk.core.document import Block, BlockType, DocumentMetadata
from chonk.loaders.base import BaseLoader, LoaderError, LoaderRegistry


@LoaderRegistry.register
class DocxLoader(BaseLoader):
    """
    Load DOCX documents using python-docx.

    Extracts:
    - Paragraphs with style information
    - Headings (from Word styles)
    - Tables with structure preserved
    - Lists
    - Document metadata
    """

    SUPPORTED_EXTENSIONS: ClassVar[list[str]] = [".docx"]
    LOADER_NAME: ClassVar[str] = "docx"

    def load(self, path: Path) -> tuple[list[Block], DocumentMetadata]:
        """Load a DOCX and extract blocks and metadata."""
        self._reset_messages()

        try:
            doc = Document(path)
            metadata = self._extract_metadata(doc, path)
            blocks = self._extract_blocks(doc)
            return blocks, metadata

        except Exception as e:
            raise LoaderError(
                f"Failed to load DOCX: {e}",
                source_path=path,
                details=str(e),
            ) from e

    def _extract_metadata(self, doc: DocxDocument, path: Path) -> DocumentMetadata:
        """Extract document metadata from DOCX."""
        core_props = doc.core_properties

        # Count words
        word_count = 0
        for para in doc.paragraphs:
            word_count += len(para.text.split())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    word_count += len(cell.text.split())

        return DocumentMetadata(
            title=core_props.title,
            author=core_props.author,
            subject=core_props.subject,
            keywords=self._parse_keywords(core_props.keywords),
            created_date=core_props.created,
            modified_date=core_props.modified,
            page_count=0,  # DOCX doesn't have fixed pages
            word_count=word_count,
            custom={
                "category": core_props.category,
                "comments": core_props.comments,
                "last_modified_by": core_props.last_modified_by,
                "revision": core_props.revision,
            },
        )

    def _parse_keywords(self, keywords_str: str | None) -> list[str]:
        """Parse keywords string into list."""
        if not keywords_str:
            return []
        # Keywords can be comma or semicolon separated
        import re

        keywords = re.split(r"[,;]", keywords_str)
        return [k.strip() for k in keywords if k.strip()]

    def _extract_blocks(self, doc: DocxDocument) -> list[Block]:
        """Extract all blocks from DOCX."""
        blocks: list[Block] = []

        # Track current list context
        current_list_id: str | None = None

        # Iterate through document body elements in order
        for element in doc.element.body:
            tag = element.tag.split("}")[-1]  # Get tag without namespace

            if tag == "p":
                # Paragraph
                para = Paragraph(element, doc)
                block = self._process_paragraph(para, current_list_id)
                if block:
                    blocks.append(block)

                    # Update list tracking
                    if block.type == BlockType.LIST_ITEM:
                        if current_list_id is None:
                            current_list_id = Block.generate_id()
                        block.parent_id = current_list_id
                    else:
                        current_list_id = None

            elif tag == "tbl":
                # Table
                table = Table(element, doc)
                block = self._process_table(table)
                if block:
                    blocks.append(block)
                current_list_id = None

        return blocks

    def _process_paragraph(self, para: Paragraph, list_id: str | None) -> Block | None:
        """Process a paragraph into a Block."""
        text = para.text.strip()
        if not text:
            return None

        # Determine block type from style
        block_type = BlockType.TEXT
        heading_level = None

        style_name = para.style.name if para.style else ""

        if style_name.startswith("Heading"):
            block_type = BlockType.HEADING
            # Extract heading level from style name (e.g., "Heading 1" -> 1)
            try:
                heading_level = int(style_name.split()[-1])
            except (ValueError, IndexError):
                heading_level = 1

        elif style_name == "Title":
            block_type = BlockType.HEADING
            heading_level = 1

        elif style_name == "Subtitle":
            block_type = BlockType.HEADING
            heading_level = 2

        elif "List" in style_name or self._is_list_paragraph(para):
            block_type = BlockType.LIST_ITEM

        elif style_name in ("Code", "Source Code", "HTML Code"):
            block_type = BlockType.CODE

        return Block(
            id=Block.generate_id(),
            type=block_type,
            content=text,
            page=1,  # DOCX doesn't have pages
            heading_level=heading_level,
            parent_id=list_id if block_type == BlockType.LIST_ITEM else None,
            metadata={
                "style": style_name,
                "is_bold": self._is_bold_paragraph(para),
                "is_italic": self._is_italic_paragraph(para),
            },
        )

    def _is_list_paragraph(self, para: Paragraph) -> bool:
        """Check if paragraph is a list item based on numbering."""
        # Check for numbering properties in the paragraph XML
        numPr = para._element.find(qn("w:numPr"))
        return numPr is not None

    def _is_bold_paragraph(self, para: Paragraph) -> bool:
        """Check if most of the paragraph is bold."""
        if not para.runs:
            return False
        bold_chars = sum(len(run.text) for run in para.runs if run.bold)
        total_chars = sum(len(run.text) for run in para.runs)
        return bold_chars > total_chars * 0.5 if total_chars > 0 else False

    def _is_italic_paragraph(self, para: Paragraph) -> bool:
        """Check if most of the paragraph is italic."""
        if not para.runs:
            return False
        italic_chars = sum(len(run.text) for run in para.runs if run.italic)
        total_chars = sum(len(run.text) for run in para.runs)
        return italic_chars > total_chars * 0.5 if total_chars > 0 else False

    def _process_table(self, table: Table) -> Block | None:
        """Process a table into a Block."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append(cells)

        if not rows:
            return None

        # Convert to markdown format
        content = self._table_to_markdown(rows)

        return Block(
            id=Block.generate_id(),
            type=BlockType.TABLE,
            content=content,
            page=1,
            metadata={
                "rows": len(rows),
                "cols": len(rows[0]) if rows else 0,
            },
        )

    def _table_to_markdown(self, rows: list[list[str]]) -> str:
        """Convert table rows to markdown format."""
        if not rows:
            return ""

        lines = []
        for i, row in enumerate(rows):
            line = "| " + " | ".join(row) + " |"
            lines.append(line)

            # Add header separator after first row
            if i == 0:
                separator = "| " + " | ".join(["---"] * len(row)) + " |"
                lines.append(separator)

        return "\n".join(lines)
